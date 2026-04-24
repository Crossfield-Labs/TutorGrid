from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from backend.knowledge.parsers.base import ParsedBlock, ParsedDocument
from backend.knowledge.parsers.image_ocr_parser import ImageOcrParser


class DocxParser:
    def __init__(self, *, enable_image_ocr: bool = True) -> None:
        self.enable_image_ocr = enable_image_ocr
        self.ocr_parser = ImageOcrParser() if enable_image_ocr else None

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for DOCX parsing.") from exc

        document = Document(str(file_path))
        blocks: list[ParsedBlock] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = str(paragraph.text or "").strip()
            if not text:
                continue
            blocks.append(
                ParsedBlock(
                    text=text,
                    section=f"paragraph-{index}",
                    page=0,
                    metadata={"paragraphIndex": index},
                )
            )
        image_texts = self._extract_document_image_texts(document)
        for image_index, image_text in enumerate(image_texts, start=1):
            blocks.append(
                ParsedBlock(
                    text=image_text,
                    section=f"image-{image_index}",
                    page=0,
                    metadata={"imageIndex": image_index, "source": "image_ocr"},
                )
            )
        return ParsedDocument(
            title=file_path.stem,
            blocks=blocks,
            metadata={"parser": "python-docx", "sourcePath": str(file_path)},
        )

    def _extract_document_image_texts(self, document: object) -> list[str]:
        if self.ocr_parser is None:
            return []
        part = getattr(document, "part", None)
        rels = getattr(part, "rels", None) if part is not None else None
        if not isinstance(rels, dict):
            return []

        image_parts: list[tuple[str, bytes]] = []
        seen_part_names: set[str] = set()
        for rel in rels.values():
            target_part = getattr(rel, "target_part", None)
            if target_part is None:
                continue
            part_name = str(getattr(target_part, "partname", "") or "")
            if "/media/" not in part_name:
                continue
            blob = getattr(target_part, "blob", b"")
            if not blob:
                continue
            if part_name in seen_part_names:
                continue
            seen_part_names.add(part_name)
            image_parts.append((part_name, bytes(blob)))

        results: list[str] = []
        with TemporaryDirectory(prefix="orchestrator_docx_ocr_") as temp_dir:
            temp_root = Path(temp_dir)
            for image_index, (part_name, blob) in enumerate(image_parts, start=1):
                ext = Path(part_name).suffix.strip(".").lower() or "png"
                image_file = temp_root / f"docx_image_{image_index}.{ext}"
                try:
                    image_file.write_bytes(blob)
                    lines = self.ocr_parser.extract_lines(image_file)
                except Exception:
                    continue
                merged = "\n".join(lines).strip()
                if merged:
                    results.append(merged)
        return results
